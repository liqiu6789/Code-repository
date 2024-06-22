from docx import Document


def replace_text_in_docx(file_path, search_text, replace_text):
    # 打开现有的Word文档
    doc = Document(file_path)

    # 遍历文档中的每一个段落
    for para in doc.paragraphs:
        if search_text in para.text:
            para.text = para.text.replace(search_text, replace_text)

    # 遍历文档中的每一个表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if search_text in cell.text:
                    cell.text = cell.text.replace(search_text, replace_text)

    # 保存修改后的文档，直接覆盖原文件
    doc.save(file_path)
    print(f"Text replaced and document saved as {file_path}")


# 使用函数
replace_text_in_docx(r"C:\Users\Administrator\Desktop\Word文档\example.docx", 'old_text', 'new_text')
