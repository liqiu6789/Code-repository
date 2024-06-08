from docx import Document

# 创建一个新的Word文档
doc = Document()

# 添加文档标题
doc.add_heading('Python自动生成的Word文档', 0)

# 添加一个段落
doc.add_paragraph('这是一个使用python-docx库生成的Word文档示例。')

# 添加一个带格式的段落
paragraph = doc.add_paragraph()
run = paragraph.add_run('这是一个加粗的段落。')
run.bold = True

# 添加一个项目符号列表
doc.add_paragraph('项目1', style='ListBullet')
doc.add_paragraph('项目2', style='ListBullet')
doc.add_paragraph('项目3', style='ListBullet')

# 添加一个编号列表
doc.add_paragraph('步骤1', style='ListNumber')
doc.add_paragraph('步骤2', style='ListNumber')
doc.add_paragraph('步骤3', style='ListNumber')

# 保存文档
doc.save(r'C:\Users\Administrator\Desktop\Word文档\example.docx')
