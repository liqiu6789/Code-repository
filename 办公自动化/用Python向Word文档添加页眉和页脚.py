from docx import Document


def add_header_footer(doc_path, header_text, footer_text):
    # 打开现有的Word文档
    doc = Document(doc_path)

    # 添加页眉
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.text = header_text

    # 添加页脚
    footer = section.footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run()
    run.text = footer_text

    # 保存文档
    doc.save(doc_path)


# 示例调用
doc_path = r"C:\Users\Administrator\Desktop\Word文档\example.docx"
header_text = '这是页眉'
footer_text = '这是页脚'
add_header_footer(doc_path, header_text, footer_text)
